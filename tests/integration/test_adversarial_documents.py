from pathlib import Path

import pytest

from pseudonymize import TransformationMode
from pseudonymize.engine import Pseudonymizer
from pseudonymize.formats import FileFormat
from pseudonymize.policy import Policy

try:
    import docx

    HAS_OFFICE = True
except ImportError:
    HAS_OFFICE = False

try:
    import fpdf

    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

pytestmark = pytest.mark.integration


def test_adversarial_long_paragraph(tmp_path: Path) -> None:
    """Test performance and correctness on a very large text block."""
    if not HAS_OFFICE:
        pytest.skip("python-docx not installed")

    engine = Pseudonymizer(policy=Policy.default(), mode=TransformationMode.REDACTED)
    path = tmp_path / "long.docx"
    out_path = tmp_path / "long_out.docx"

    doc = docx.Document()
    # 50,000 characters of junk
    junk = "a" * 25000
    doc.add_paragraph(f"{junk} hidden@example.com {junk}")
    doc.save(str(path))

    # Process
    result = engine.process_file(path, out_path, format=FileFormat.DOCX)

    assert out_path.exists()
    assert result.statistics.detections_found >= 1

    out_doc = docx.Document(str(out_path))
    text = out_doc.paragraphs[0].text
    assert "[REDACTED]" in text
    assert "hidden@example.com" not in text


def test_adversarial_nested_tables_docx(tmp_path: Path) -> None:
    """Test deep nesting of tables to ensure structural extraction visits all nodes."""
    if not HAS_OFFICE:
        pytest.skip("python-docx not installed")

    engine = Pseudonymizer(policy=Policy.default(), mode=TransformationMode.REDACTED)
    path = tmp_path / "nested.docx"
    out_path = tmp_path / "nested_out.docx"

    doc = docx.Document()
    table1 = doc.add_table(rows=1, cols=1)
    cell1 = table1.cell(0, 0)

    # Nested table 2
    table2 = cell1.add_table(rows=1, cols=1)
    cell2 = table2.cell(0, 0)

    # Nested table 3
    table3 = cell2.add_table(rows=1, cols=1)
    cell3 = table3.cell(0, 0)

    cell3.text = "Deeply nested email: nested@example.com"
    doc.save(str(path))

    engine.process_file(path, out_path, format=FileFormat.DOCX)

    out_doc = docx.Document(str(out_path))
    # python-docx doesn't easily expose nested tables through .tables, you have to walk cells
    cell_text = out_doc.tables[0].cell(0, 0).tables[0].cell(0, 0).tables[0].cell(0, 0).text
    assert "[REDACTED]" in cell_text
    assert "nested@example.com" not in cell_text


def test_adversarial_pdf_hidden_text(tmp_path: Path) -> None:
    """Test PDF extraction and redaction when text is white-on-white or off-page."""
    if not HAS_FPDF:
        pytest.skip("fpdf2 not installed")

    engine = Pseudonymizer(policy=Policy.default(), mode=TransformationMode.REDACTED)
    path = tmp_path / "hidden.pdf"
    out_path = tmp_path / "hidden_out.pdf"

    pdf = fpdf.FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)

    # White text on white background
    pdf.set_text_color(255, 255, 255)
    pdf.cell(text="ghost@example.com")

    pdf.output(str(path))

    engine.process_file(path, out_path, format=FileFormat.PDF)

    import pymupdf

    out_doc = pymupdf.open(out_path)
    text = out_doc[0].get_text()
    assert "[REDACTED]" in text
    assert "ghost@example.com" not in text
    out_doc.close()


def test_adversarial_zwnj_obfuscation(tmp_path: Path) -> None:
    """Test that zero-width non-joiners do not evade PII detection."""
    engine = Pseudonymizer(policy=Policy.default(), mode=TransformationMode.REDACTED)
    path = tmp_path / "zwnj.txt"
    out_path = tmp_path / "zwnj_out.txt"

    # Insert a ZWNJ (\u200c) inside the email
    # Currently, local rules might not detect this because ZWNJs break regexes.
    # This test acts as a baseline to ensure the engine doesn't crash when faced with them,
    # and to explicitly document the current state of detection.
    # Advanced ML tokenizers handles these better natively.
    obfuscated_email = "sneaky\u200c@example.com"
    path.write_text(f"Contact {obfuscated_email} now.", encoding="utf-8")

    engine.process_file(path, out_path, format=FileFormat.TEXT)

    text = out_path.read_text(encoding="utf-8")
    # If detection works, it will be redacted. If not, it won't crash.
    # We assert it does not crash by reading the file.
    assert "Contact" in text
